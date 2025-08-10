"""
FIXED main.py - This version definitely works!
Save as: backend/main.py
"""

from fastapi import FastAPI, File, UploadFile, BackgroundTasks, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
import uuid
from pathlib import Path
from typing import List
import zipfile
import logging
import sys
import os

#from doc_processor.latex_processor  import process_word_document, save_results
from doc_processor.main_word_com_equation_replacer import WordCOMEquationReplacer


# Setup simple logging to console
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('processing.log',encoding='utf-8')
    ]
)
logger = logging.getLogger(__name__)

app = FastAPI(title="Document Processing API")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Fix paths - use absolute paths to avoid confusion
BASE_DIR = Path(__file__).parent.absolute()
TEMP_DIR = BASE_DIR / "temp"
OUTPUT_DIR = BASE_DIR / "output"

# Create directories
TEMP_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

logger.info(f"Base directory: {BASE_DIR}")
logger.info(f"Temp directory: {TEMP_DIR}")
logger.info(f"Output directory: {OUTPUT_DIR}")

# In-memory job tracking
jobs = {}

@app.post("/api/process")
async def process_documents(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    processor_type: str = Form("word_to_html")
):
    """Process documents with specified processor"""
    job_id = str(uuid.uuid4())
    
    logger.info(f"=== NEW JOB {job_id} ===")
    logger.info(f"Processing {len(files)} files with {processor_type}")
    
    # Initialize job
    jobs[job_id] = {
        "status": "processing",
        "total": len(files),
        "completed": 0,
        "results": [],
        "processor": processor_type
    }
    
    # Create job directories
    job_temp_dir = TEMP_DIR / job_id
    job_output_dir = OUTPUT_DIR / job_id
    job_temp_dir.mkdir(parents=True, exist_ok=True)
    job_output_dir.mkdir(parents=True, exist_ok=True)
    
    logger.info(f"Job temp dir: {job_temp_dir}")
    logger.info(f"Job output dir: {job_output_dir}")
    
    # Save uploaded files
    file_paths = []
    for file in files:
        file_path = job_temp_dir / file.filename
        content = await file.read()
        file_path.write_bytes(content)
        file_paths.append(file_path)
        logger.info(f"Saved uploaded file: {file_path}")
    
    # Process in background
    background_tasks.add_task(
        process_job, 
        job_id, 
        file_paths, 
        processor_type,
        job_output_dir
    )
    
    return {
        "job_id": job_id,
        "status": "processing",
        "message": f"Processing {len(files)} documents"
    }

@app.get("/api/status/{job_id}")
async def get_status(job_id: str):
    """Check processing status"""
    if job_id not in jobs:
        logger.error(f"Job {job_id} not found")
        return JSONResponse(status_code=404, content={"error": "Job not found"})
    
    job = jobs[job_id]
    return {
        "job_id": job_id,
        "status": job["status"],
        "progress": f"{job['completed']}/{job['total']}",
        "processor": job["processor"],
        "results": job.get("results", [])
    }

"""
Add this to your main.py - Replace the download endpoints with these
"""

@app.get("/api/download/{job_id}")
async def download_all_results(job_id: str):
    """Download all processed results as ZIP or single file"""
    logger.info(f"Download all request for job {job_id}")
    
    if job_id not in jobs:
        logger.error(f"Job {job_id} not found")
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = jobs[job_id]
    if job["status"] != "completed":
        logger.error(f"Job {job_id} not completed yet")
        raise HTTPException(status_code=400, detail="Job not completed")
    
    results = job["results"]
    successful_results = [r for r in results if "error" not in r]
    
    if not successful_results:
        logger.error("No successful results")
        raise HTTPException(status_code=404, detail="No successful results")
    
    if len(results) == 1 and results[0].get("type") == "application/zip":
        # Direct ZIP download
        file_path = Path(results[0]["path"])
        logger.info(f"Serving ZIP file: {file_path}")
        return FileResponse(
            path=str(file_path),
            filename=results[0]["output_filename"],
            media_type="application/zip"
        )

    if len(successful_results) == 1:
        # Single file - return directly
        file_path = Path(successful_results[0]["path"])
        logger.info(f"Downloading single file: {file_path}")
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            raise HTTPException(status_code=404, detail=f"File not found")
        
        return FileResponse(
            path=str(file_path),
            filename=successful_results[0].get("output_filename", successful_results[0]["filename"])
        )
    else:
        # Multiple files - create ZIP
        zip_path = OUTPUT_DIR / job_id / f"results_{job_id}.zip"
        logger.info(f"Creating ZIP: {zip_path}")
        
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for result in successful_results:
                file_path = Path(result["path"])
                if file_path.exists():
                    output_name = result.get("output_filename", result["filename"])
                    zipf.write(file_path, output_name)
                    logger.info(f"Added to ZIP: {output_name}")
        
        return FileResponse(
            path=str(zip_path),
            filename=f"results_{job_id}.zip"
        )

@app.get("/api/download/{job_id}/{index}")
async def download_single_result(job_id: str, index: int):
    """Download a single result file by index"""
    logger.info(f"Download single file request: job={job_id}, index={index}")
    
    if job_id not in jobs:
        logger.error(f"Job {job_id} not found")
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = jobs[job_id]
    
    # Filter successful results only
    successful_results = [r for r in job["results"] if "error" not in r]
    
    if index >= len(successful_results):
        logger.error(f"Invalid file index: {index}")
        raise HTTPException(status_code=404, detail="Invalid file index")
    
    result = successful_results[index]
    file_path = Path(result["path"])
    
    if not file_path.exists():
        logger.error(f"File not found: {file_path}")
        raise HTTPException(status_code=404, detail=f"File not found")
    
    logger.info(f"Serving file: {file_path}")

    media_type = result.get("type", "application/octet-stream")  # ADD THIS LINE
    
    return FileResponse(
        path=str(file_path),
        filename=result.get("output_filename", result["filename"]),
        media_type=media_type  
    )


def should_zip_output(output_dir):
    """
    Determine if output should be zipped
    Returns True if:
    - More than 2 files in output
    - Any subdirectories exist
    """
    output_path = Path(output_dir)
    
    # Get all files and dirs
    all_items = list(output_path.iterdir())
    files = [f for f in all_items if f.is_file()]
    dirs = [d for d in all_items if d.is_dir()]
    is_zipp_output = len(dirs) > 0 or len(files) > 2
    # Zip if subdirectories exist or more than 2 files
    print (f"is_zipp_output = {is_zipp_output}")
    return is_zipp_output

def create_zip_output(output_dir, job_id):
    """
    Create a zip file of the output directory
    """
    output_path = Path(output_dir)
    zip_filename = f"{job_id}_output.zip"
    zip_path = output_path / zip_filename
    
    print(f"\n📦 Creating zip file: {zip_filename}")
    
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        # Add all files and folders
        for item in output_path.iterdir():
            if item.name != zip_filename:  # Don't include the zip itself
                if item.is_file():
                    zipf.write(item, item.name)
                    print(f"  Added file: {item.name}")
                elif item.is_dir():
                    # Add directory and its contents
                    for root, dirs, files in os.walk(item):
                        root_path = Path(root)
                        for file in files:
                            file_path = root_path / file
                            arcname = file_path.relative_to(output_path)
                            zipf.write(file_path, arcname)
                            print(f"  Added: {arcname}")
    
    print(f"✅ Zip created: {zip_path}")
    return zip_path


async def process_job(job_id: str, file_paths: List[Path], processor_type: str, output_dir: Path):
    """Background job processor with fixed ZIP handling"""
    logger.info(f"Starting background processing for job {job_id}")
    logger.info(f"Output directory: {output_dir}")
    
    # Temporary list for results
    temp_results = []
    
    for i, file_path in enumerate(file_paths):
        try:
            logger.info(f"Processing file {i+1}/{len(file_paths)}: {file_path.name}")
            
            if processor_type == "word_to_html":
                from full_word_processor.WordFullProcessor import WordFullProcessor
                processor = WordFullProcessor()
                output_file = processor.process_document(str(file_path), str(output_dir))
            
            elif processor_type == "latex_equations":
                replacer = WordCOMEquationReplacer()
                output_filename = f"{Path(file_path).stem}_latex_equations.docx"
                output_path = os.path.join(output_dir, output_filename)
                output_file = replacer.process_document(file_path, output_path)
                if output_file:
                    output_file = Path(output_file)
                    
            elif processor_type == "word_complete":
                from full_word_processor.WordFullProcessor import WordFullProcessor
                
                # Step 1: Process equations
                replacer = WordCOMEquationReplacer()
                equations_filename = f"{Path(file_path).stem}_equations.docx"
                equations_path = os.path.join(output_dir, equations_filename)
                equations_doc = replacer.process_document(file_path, equations_path)
                
                # Step 2: Convert to HTML
                if equations_doc:
                    processor = WordFullProcessor()
                    output_file = processor.process_document(equations_doc, output_dir)
                else:
                    processor = WordFullProcessor()
                    output_file = processor.process_document(file_path, output_dir)
            
            else:  # scan_verify
                output_file = await scan_and_verify(file_path, output_dir)
            
            result = {
                "filename": file_path.name,
                "output_filename": output_file.name,
                "path": str(output_file),
                "index": i,
                "success": True
            }
            
            temp_results.append(result)
            jobs[job_id]["completed"] += 1
            
            logger.info(f"Successfully processed: {file_path.name}")
            logger.info(f"  Output: {output_file}")
            
        except Exception as e:
            logger.error(f"Failed to process {file_path.name}: {str(e)}")
            temp_results.append({
                "filename": file_path.name,
                "error": str(e),
                "index": i
            })
            jobs[job_id]["completed"] += 1
    
    # THIS MUST BE AT THE SAME INDENTATION AS THE FOR LOOP
    logger.info(f"DEBUG: Processing complete, checking if should zip...")
    logger.info(f"DEBUG: temp_results count: {len(temp_results)}")
    logger.info(f"DEBUG: processor_type: {processor_type}")
    
    if processor_type in ["word_complete", "word_to_html"]:
        # Check output directory contents
        output_files = list(output_dir.iterdir())
        logger.info(f"DEBUG: Output dir contains {len(output_files)} items")
        for f in output_files:
            logger.info(f"  - {f.name} ({'dir' if f.is_dir() else 'file'})")
        
        # Simplified ZIP logic - always ZIP if word_complete with multiple files
        # or if there are subdirectories
        has_subdirs = any(f.is_dir() for f in output_files)
        has_multiple_files = len([f for f in output_files if f.is_file()]) > 2
        
        should_zip = has_subdirs or has_multiple_files or processor_type == "word_complete"
        
        logger.info(f"DEBUG: has_subdirs={has_subdirs}, has_multiple_files={has_multiple_files}")
        logger.info(f"DEBUG: should_zip={should_zip}")
        
        if should_zip:
            logger.info(f"DEBUG: Creating ZIP file")
            zip_file = create_zip_output(output_dir, job_id)
            
            logger.info(f"DEBUG: ZIP created at {zip_file}")
            
            zip_result = {
                "filename": zip_file.name,
                "output_filename": zip_file.name,
                "path": str(zip_file),
                "index": 0,
                "success": True,
                "size": zip_file.stat().st_size,
                "type": "application/zip"
            }
            
            # Replace results with just the ZIP
            jobs[job_id]["results"] = [zip_result]
            logger.info(f"✅ ZIP REPLACEMENT DONE: {jobs[job_id]['results']}")
        else:
            logger.info(f"DEBUG: Not zipping, using individual files")
            jobs[job_id]["results"] = temp_results
    else:
        jobs[job_id]["results"] = temp_results
    
    jobs[job_id]["status"] = "completed"
    logger.info(f"=== JOB {job_id} COMPLETED ===")
    logger.info(f"=== FINAL RESULTS: {len(jobs[job_id]['results'])} items ===")

async def process_job_old(job_id: str, file_paths: List[Path], processor_type: str, output_dir: Path):
    """Background job processor"""
    logger.info(f"Starting background processing for job {job_id}")
    logger.info(f"Output directory: {output_dir}")
    
    # ONLY ADD THIS LINE - temporary results during processing
    temp_results = []
    
    for i, file_path in enumerate(file_paths):
        try:
            logger.info(f"Processing file {i+1}/{len(file_paths)}: {file_path.name}")
            
            '''
            if processor_type == "word_to_html":
                # Simple HTML conversion
                output_file = await convert_to_html(file_path, output_dir)
            else:  # scan_verify
                output_file = await scan_and_verify(file_path, output_dir)
            if processor_type == "word_to_html":
                # Simple HTML conversion
                output_file = await convert_to_html(file_path, output_dir)
            '''
            
            if processor_type == "word_to_html":
                # Simple HTML conversion
                output_file = await convert_to_html(file_path, output_dir)
            
            elif processor_type == "latex_equations":
                # Just call your replacer directly - no await, no asyncio!
                replacer = WordCOMEquationReplacer()
                output_filename = f"{Path(file_path).stem}_latex_equations.docx"
                output_path = os.path.join(output_dir, output_filename)
                
                # Direct synchronous call - exactly like your existing code does!
                output_file = replacer.process_document(file_path, output_path)
                
                # Convert to Path for consistency with other processors
                if output_file:
                    output_file = Path(output_file)
            elif processor_type == "word_complete":
                #from full-word-processor.word_com_equation_replacer import WordCOMEquationReplacer
                from full_word_processor.WordFullProcessor import WordFullProcessor
                
                
                # Step 1: Process equations
                replacer = WordCOMEquationReplacer()
                equations_filename = f"{Path(file_path).stem}_equations.docx"
                equations_path = os.path.join(output_dir, equations_filename)
                
                equations_doc = replacer.process_document(file_path, equations_path)
                
                # Step 2: Convert to HTML
                if equations_doc:
                    processor = WordFullProcessor()
                    output_file = processor.process_document(equations_doc, output_dir)
                else:
                    # If equation processing failed, use original
                    processor = WordFullProcessor()
                    output_file = processor.process_document(file_path, output_dir)
            
            else:  # scan_verify
                output_file = await scan_and_verify(file_path, output_dir)
            
            result = {
                "filename": file_path.name,
                "output_filename": output_file.name,
                "path": str(output_file),
                "index": i,
                "success": True
            }
            
            temp_results.append(result)  # CHANGE: temp_results instead of jobs[job_id]["results"]
            jobs[job_id]["completed"] += 1
            
            logger.info(f"Successfully processed: {file_path.name}")
            logger.info(f"  Output: {output_file}")
            
        except Exception as e:
            logger.error(f"Failed to process {file_path.name}: {str(e)}")
            temp_results.append({  # CHANGE: temp_results instead of jobs[job_id]["results"]
                "filename": file_path.name,
                "error": str(e),
                "index": i
            })
            jobs[job_id]["completed"] += 1
    
    logger.info(f"DEBUG: temp_results has {len(temp_results)} items")
    logger.info(f"DEBUG: processor_type = {processor_type}")

    # After all processing, check if we should zip
    if processor_type in ["word_complete", "word_to_html"]:
        logger.info(f"DEBUG: Checking if should zip...")

        if should_zip_output(output_dir):
            logger.info(f"DEBUG: Should zip = True")

            # Create zip file
            zip_file = create_zip_output(output_dir, job_id)
            logger.info(f"DEBUG: Zip created at {zip_file}")

            # Clear results and add only zip
            '''
            results = [{
                "filename": zip_file.name,
                "size": zip_file.stat().st_size,
                "type": "application/zip"
            }]
            print(f"\n📦 Output zipped due to multiple files/folders")
            '''

            # Create new results with ONLY the ZIP
            zip_result = {
                "filename": zip_file.name,
                "output_filename": zip_file.name,
                "path": str(zip_file),
                "index": 0,
                "success": True,
                "size": zip_file.stat().st_size,
                "type": "application/zip"
            }
            
            # REPLACE all results with just the ZIP
            logger.info(f"DEBUG: Before replace - jobs[{job_id}]['results'] = {jobs[job_id].get('results', [])}")
            jobs[job_id]["results"] = [zip_result]  # Direct assignment
            logger.info(f"DEBUG: After replace - jobs[{job_id}]['results'] = {jobs[job_id]['results']}")
            logger.info(f"📦 Results replaced with ZIP: {jobs[job_id]['results'][0]['filename']}")
            
            logger.info(f"📦 Results replaced with ZIP: {jobs[job_id]['results']}")

        else:
            # Add individual files to results
            logger.info(f"DEBUG: Should zip = False")
            jobs[job_id]["results"] = temp_results  # ADD: Set results from temp
            for file in output_dir.iterdir():
                if file.is_file():
                    jobs[job_id]["results"].append({  # FIX: Changed 'results' to 'jobs[job_id]["results"]'
                        "filename": file.name,
                        "size": file.stat().st_size,
                        "type": "text/html" if file.suffix == ".html" else "application/octet-stream"
                    })
    else:
        # For other processors, list files normally
        logger.info(f"DEBUG: Other processor type: {processor_type}")

        jobs[job_id]["results"] = temp_results  # ADD: Set results from temp
        for file in output_dir.iterdir():
            if file.is_file():
                jobs[job_id]["results"].append({  # FIX: Changed 'results' to 'jobs[job_id]["results"]'
                    "filename": file.name,
                    "size": file.stat().st_size,
                    "type": "application/octet-stream"
                })
        

    logger.info(f"DEBUG: Final results count: {len(jobs[job_id]['results'])}")
    logger.info(f"DEBUG: Final results: {jobs[job_id]['results']}")

    jobs[job_id]["status"] = "completed"
    logger.info(f"=== JOB {job_id} COMPLETED ===")


async def convert_to_html(input_file: Path, output_dir: Path) -> Path:
    """Simple HTML conversion using mammoth"""
    import mammoth
    
    logger.info(f"Converting to HTML: {input_file.name}")
    
    # Convert with mammoth
    with open(input_file, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
    
    # Create HTML with Arabic support
    html_content = f"""<!DOCTYPE html>
<html lang="ar" xx="1" dir="rtl">
<head>
    <meta charset="UTF-8">
    <title>{input_file.stem}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            direction: rtl;
        }}
        h1, h2, h3 {{ color: #333; }}
        img {{ max-width: 100%; height: auto; }}
    </style>
        <script>
      window.MathJax = {
        tex: {
          inlineMath: [['\\(', '\\)']],
          displayMath: [['\\[', '\\]']]
        }
      };
    </script>
    <script src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-chtml.js"></script>

</head>
<body>
{result.value}
</body>
</html>"""
    
    # Save HTML
    output_file = output_dir / f"{input_file.stem}.html"
    output_file.write_text(html_content, encoding='utf-8')
    
    logger.info(f"HTML saved to: {output_file}")
    return output_file

async def scan_and_verify(input_file: Path, output_dir: Path) -> Path:
    """Simple document analysis"""
    from docx import Document
    import pandas as pd
    
    logger.info(f"Scanning document: {input_file.name}")
    
    doc = Document(input_file)
    
    # Basic analysis
    analysis = {
        "Filename": input_file.name,
        "Paragraphs": len(doc.paragraphs),
        "Tables": len(doc.tables),
        "Word Count": sum(len(p.text.split()) for p in doc.paragraphs if p.text)
    }
    
    # Save to Excel
    output_file = output_dir / f"{input_file.stem}_analysis.xlsx"
    df = pd.DataFrame([analysis])
    df.to_excel(output_file, index=False)
    
    logger.info(f"Analysis saved to: {output_file}")
    return output_file

@app.get("/")
async def root():
    """Health check"""
    return {
        "status": "running",
        "message": "Document Processing API",
        "temp_dir": str(TEMP_DIR),
        "output_dir": str(OUTPUT_DIR)
    }

@app.get("/api/debug/{job_id}")
async def debug_job(job_id: str):
    """Debug endpoint to check job details"""
    if job_id not in jobs:
        return {"error": "Job not found"}
    
    job = jobs[job_id]
    job_output_dir = OUTPUT_DIR / job_id
    
    # List files in output directory
    output_files = []
    if job_output_dir.exists():
        output_files = [f.name for f in job_output_dir.iterdir()]
    
    return {
        "job": job,
        "output_dir": str(job_output_dir),
        "output_files": output_files,
        "output_dir_exists": job_output_dir.exists()
    }

if __name__ == "__main__":
    import uvicorn
    logger.info("Starting Document Processing API...")
    logger.info(f"Working directory: {Path.cwd()}")
    logger.info(f"Script location: {Path(__file__).parent}")
    uvicorn.run(app, host="0.0.0.0", port=8000)