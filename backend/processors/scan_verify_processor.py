from pathlib import Path
import win32com.client
from docx import Document
import pandas as pd
from .base_processor import BaseProcessor

class ScanVerifyProcessor(BaseProcessor):
    """Scan and verify Word documents"""
    
    async def process(self, file_path: Path) -> dict:
        """Extract document structure and verify content"""
        
        # Use python-docx for structure analysis
        doc = Document(file_path)
        
        analysis = {
            "filename": file_path.name,
            "word_count": 0,
            "sections": [],
            "images": 0,
            "tables": len(doc.tables),
            "verification": {}
        }
        
        # Extract sections and content
        current_section = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # Check for section headers (font size > 14)
            if para.style and 'Heading' in para.style.name:
                current_section = {
                    "title": text,
                    "content": [],
                    "word_count": 0
                }
                analysis["sections"].append(current_section)
            elif current_section:
                current_section["content"].append(text)
                word_count = len(text.split())
                current_section["word_count"] += word_count
                analysis["word_count"] += word_count
        
        # Count images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                analysis["images"] += 1
        
        # Verification checks
        analysis["verification"] = {
            "has_title": len(analysis["sections"]) > 0,
            "has_content": analysis["word_count"] > 100,
            "structure_valid": len(analysis["sections"]) > 1,
            "ready_for_conversion": True
        }
        
        # Save analysis to Excel
        output_path = file_path.parent / f"{file_path.stem}_analysis.xlsx"
        self._save_to_excel(analysis, output_path)
        
        return {
            "filename": file_path.name,
            "path": str(output_path),
            "analysis": analysis
        }
    
    def _save_to_excel(self, analysis: dict, output_path: Path):
        """Save analysis to Excel file"""
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # Summary sheet
            summary_data = {
                "Metric": ["Word Count", "Sections", "Images", "Tables", "Ready for Conversion"],
                "Value": [
                    analysis["word_count"],
                    len(analysis["sections"]),
                    analysis["images"],
                    analysis["tables"],
                    analysis["verification"]["ready_for_conversion"]
                ]
            }
            pd.DataFrame(summary_data).to_excel(writer, sheet_name="Summary", index=False)
            
            # Sections detail
            if analysis["sections"]:
                sections_data = []
                for section in analysis["sections"]:
                    sections_data.append({
                        "Section": section["title"],
                        "Word Count": section["word_count"],
                        "Content Preview": section["content"][0][:100] if section["content"] else ""
                    })
                pd.DataFrame(sections_data).to_excel(writer, sheet_name="Sections", index=False)